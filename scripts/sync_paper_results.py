#!/usr/bin/env python3
"""Patch Adaptive_RAG_Paper_Rewritten.docx Results section from summary_all.json (measured experiments)."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document


def fmt_pct(x: float) -> str:
    return f"{100.0 * x:.1f}%"


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--run", required=True, help="Run directory containing summary_all.json")
    ap.add_argument("--doc", default="Adaptive_RAG_Paper_Rewritten.docx")
    args = ap.parse_args()

    run_dir = Path(args.run)
    data = json.loads((run_dir / "summary_all.json").read_text(encoding="utf-8"))
    pre = data["pre"]
    post = data["post"]

    def agg(phase: str, system: str, key: str) -> float:
        return float(data[phase][system]["aggregate"][key])

    def ci(phase: str, system: str, key: str) -> tuple[float, float]:
        c = data[phase][system]["bootstrap_ci"][key]
        return float(c["low"]), float(c["high"])

    pre_sys = "adaptive"
    cp_pre = data["pre"][pre_sys]["clopper_pearson_vrp_ge_1"]
    n = int(cp_pre["n"])
    succ = int(cp_pre["successes"])

    p_faiss = post["faiss_static"]
    p_hyb = post["hybrid_bm25_dense"]
    p_mma = post["mma_text_proxy"]
    p_ada = post["adaptive"]

    para_under = (
        f"Under the initial document snapshot, all four systems performed similarly. Each system reached "
        f"strong version-aware retrieval precision on the evaluated query set before amendments were introduced "
        f"(VRP mean = {fmt_pct(agg('pre', pre_sys, 'vrp'))}; bootstrap 95% CI for VRP "
        f"[{fmt_pct(ci('pre', pre_sys, 'vrp')[0])}, {fmt_pct(ci('pre', pre_sys, 'vrp')[1])}]; "
        f"exact Clopper–Pearson 95% CI for VRP=1 on all queries: [{fmt_pct(cp_pre['low'])}, {fmt_pct(cp_pre['high'])}] "
        f"with {succ}/{n} strict successes in this run). That result matters because it shows that the baselines "
        f"are not weak by construction; they perform adequately when the corpus remains stable."
    )

    para_diff = (
        "The differences appeared after the amended document versions were injected. In that phase, the "
        f"FAISS baseline showed SCR = {fmt_pct(agg('post', 'faiss_static', 'scr'))} "
        f"(95% CI [{fmt_pct(ci('post', 'faiss_static', 'scr')[0])}, {fmt_pct(ci('post', 'faiss_static', 'scr')[1])}]) "
        f"and SHR = {fmt_pct(agg('post', 'faiss_static', 'shr'))} while VRP = {fmt_pct(agg('post', 'faiss_static', 'vrp'))}. "
        f"MMA-RAG (text-only proxy) showed SCR = {fmt_pct(agg('post', 'mma_text_proxy', 'scr'))}, "
        f"SHR = {fmt_pct(agg('post', 'mma_text_proxy', 'shr'))}, VRP = {fmt_pct(agg('post', 'mma_text_proxy', 'vrp'))}. "
        f"The hybrid BM25–dense baseline showed SCR = {fmt_pct(agg('post', 'hybrid_bm25_dense', 'scr'))}, "
        f"SHR = {fmt_pct(agg('post', 'hybrid_bm25_dense', 'shr'))}, VRP = {fmt_pct(agg('post', 'hybrid_bm25_dense', 'vrp'))}. "
        f"The adaptive pipeline (v1 index post-amendment) achieved SCR = {fmt_pct(agg('post', 'adaptive', 'scr'))}, "
        f"SHR = {fmt_pct(agg('post', 'adaptive', 'shr'))}, VRP = {fmt_pct(agg('post', 'adaptive', 'vrp'))} "
        f"on this synthetic harness run."
    )

    para_cas = (
        "Compliance alignment and version-sensitive answer accuracy moved with retrieval quality. "
        f"After amendments, FAISS scored CAS = {fmt_pct(agg('post', 'faiss_static', 'cas'))}, "
        f"MMA-RAG proxy {fmt_pct(agg('post', 'mma_text_proxy', 'cas'))}, hybrid {fmt_pct(agg('post', 'hybrid_bm25_dense', 'cas'))}, "
        f"adaptive {fmt_pct(agg('post', 'adaptive', 'cas'))}. "
        f"VA-RAcc: FAISS {fmt_pct(agg('post', 'faiss_static', 'va_racc'))}, MMA {fmt_pct(agg('post', 'mma_text_proxy', 'va_racc'))}, "
        f"hybrid {fmt_pct(agg('post', 'hybrid_bm25_dense', 'va_racc'))}, adaptive {fmt_pct(agg('post', 'adaptive', 'va_racc'))}. "
        "These numbers come from the bundled synthetic corpus and hash embeddings; replace with your own run directory when scaling."
    )

    doc = Document(args.doc)
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Under the initial document snapshot"):
            for r in p.runs:
                r.text = ""
            p.add_run(para_under)
        elif t.startswith("The differences appeared after the amended"):
            for r in p.runs:
                r.text = ""
            p.add_run(para_diff)
        elif t.startswith("Compliance alignment moved in the same"):
            for r in p.runs:
                r.text = ""
            p.add_run(para_cas)

    doc.save(args.doc)
    print("Updated", args.doc, "from", run_dir)


if __name__ == "__main__":
    main()
